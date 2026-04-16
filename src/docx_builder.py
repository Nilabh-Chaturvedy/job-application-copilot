from docx import Document


def build_docx_resume(structured_resume: dict, file_path: str):
    doc = Document()

    # ---- SUMMARY ----
    summary = structured_resume.get("summary", "")
    if summary:
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(summary)

    # ---- EXPERIENCE ----
    experience = structured_resume.get("experience", [])
    if experience:
        doc.add_heading("Experience", level=1)

        for exp in experience:
            role = exp.get("role", "")
            company = exp.get("company", "")
            location = exp.get("location", "")
            dates = exp.get("dates", "")

            header = f"{role} - {company}"
            doc.add_paragraph(header, style="Heading 2")

            if location or dates:
                doc.add_paragraph(f"{location} | {dates}")

            for bullet in exp.get("bullets", []):
                doc.add_paragraph(bullet, style="List Bullet")

    # ---- SKILLS ----
    skills = structured_resume.get("skills", [])
    if skills:
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(skills))

    # ---- EDUCATION ----
    education = structured_resume.get("education", [])
    if education:
        doc.add_heading("Education", level=1)

        for edu in education:
            line = f"{edu.get('degree', '')} - {edu.get('institution', '')} ({edu.get('dates', '')})"
            doc.add_paragraph(line)

    doc.save(file_path)